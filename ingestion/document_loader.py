"""
Chargeur de documents (PDF, TXT, DOCX)
"""
import re
from pathlib import Path
from typing import List, Dict, Optional
from pypdf import PdfReader
from docx import Document
from utils.logger import log


class DocumentLoader:
    """Charge differents types de documents"""
    
    SUPPORTED_EXTENSIONS = ['.pdf', '.txt', '.docx', '.md']
    
    def __init__(self):
        """Initialise le loader"""
        log.info("DocumentLoader initialise")
    
    def load(self, file_path: str) -> Dict[str, any]:
        """
        Charge un document
        
        Args:
            file_path: Chemin vers le fichier
            
        Returns:
            Dict avec 'text', 'metadata', 'source'
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Fichier introuvable: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Extension non supportee: {extension}")
        
        log.info(f"Chargement document: {path.name}")
        
        if extension == '.pdf':
            text = self._load_pdf(path)
        elif extension == '.txt' or extension == '.md':
            text = self._load_text(path)
        elif extension == '.docx':
            text = self._load_docx(path)
        else:
            raise ValueError(f"Extension non supportee: {extension}")
        
        text = self._clean_text(text)
        
        metadata = {
            'filename': path.name,
            'extension': extension,
            'size_bytes': path.stat().st_size,
            'num_chars': len(text)
        }
        
        log.info(f"Document charge: {len(text)} caracteres")
        
        return {
            'text': text,
            'metadata': metadata,
            'source': str(path)
        }
    
    def _load_pdf(self, path: Path) -> str:
        """Charge un PDF"""
        try:
            reader = PdfReader(str(path))
            text = ""
            
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                text += page_text + "\n\n"
            
            log.debug(f"PDF charge: {len(reader.pages)} pages")
            return text
            
        except Exception as e:
            log.error(f"Erreur chargement PDF: {e}")
            raise
    
    def _load_text(self, path: Path) -> str:
        """Charge un fichier texte"""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            log.debug(f"Fichier texte charge: {len(text)} caracteres")
            return text
            
        except UnicodeDecodeError:
            try:
                with open(path, 'r', encoding='latin-1') as f:
                    text = f.read()
                log.debug("Fichier charge avec encoding latin-1")
                return text
            except Exception as e:
                log.error(f"Erreur chargement texte: {e}")
                raise
    
    def _load_docx(self, path: Path) -> str:
        """Charge un fichier Word"""
        try:
            doc = Document(str(path))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            log.debug(f"DOCX charge: {len(doc.paragraphs)} paragraphes")
            return text
            
        except Exception as e:
            log.error(f"Erreur chargement DOCX: {e}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Nettoie le texte"""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = text.strip()
        return text
    
    def load_directory(self, directory: str) -> List[Dict]:
        """
        Charge tous les documents d'un dossier
        
        Args:
            directory: Chemin vers le dossier
            
        Returns:
            Liste de documents charges
        """
        dir_path = Path(directory)
        
        if not dir_path.exists():
            raise FileNotFoundError(f"Dossier introuvable: {directory}")
        
        documents = []
        
        for ext in self.SUPPORTED_EXTENSIONS:
            for file_path in dir_path.glob(f"**/*{ext}"):
                try:
                    doc = self.load(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    log.error(f"Erreur chargement {file_path}: {e}")
        
        log.info(f"Dossier charge: {len(documents)} documents")
        return documents


document_loader = DocumentLoader()

